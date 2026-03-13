from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def add_heading_with_color(doc, text, level=1):
    """Add a heading with custom color"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(24, 72, 160)  # #1848A0
    return heading

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# ==================== TITLE PAGE ====================
title = doc.add_heading('HealthX', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(48)
    run.font.color.rgb = RGBColor(24, 72, 160)

subtitle = doc.add_paragraph('IoT-Based Health Monitoring System')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.size = Pt(20)
    run.font.italic = True

doc.add_paragraph()
system_doc = doc.add_paragraph('System Documentation')
system_doc.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in system_doc.runs:
    run.font.size = Pt(16)
    run.font.bold = True

doc.add_paragraph()
version_info = doc.add_paragraph(f'Version 1.0\nGenerated: {datetime.now().strftime("%B %d, %Y")}')
version_info.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ==================== TABLE OF CONTENTS ====================
add_heading_with_color(doc, 'Table of Contents', 1)
toc_items = [
    '1. System Overview',
    '2. System Architecture', 
    '3. Technology Stack',
    '4. Project Structure',
    '5. Core Components',
    '   5.1 API Configuration',
    '   5.2 API Service',
    '   5.3 ESP32 Connection Service',
    '   5.4 Data Models',
    '6. API Endpoints',
    '   6.1 Web Server API',
    '   6.2 ESP32 Sensor API',
    '7. User Interfaces',
    '   7.1 Authentication Screens',
    '   7.2 Health Worker Dashboard',
    '   7.3 User Dashboard',
    '8. Data Flow',
    '9. Installation Guide',
    '10. Configuration',
    '11. User Guide',
    '12. Troubleshooting',
    '13. Future Enhancements'
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number' if not item.startswith('   ') else 'Normal')

doc.add_page_break()

# ==================== 1. SYSTEM OVERVIEW ====================
add_heading_with_color(doc, '1. System Overview', 1)

doc.add_paragraph(
    'HealthX is a comprehensive IoT-based health monitoring system designed to provide real-time '
    'health data collection and monitoring capabilities through integration of ESP32-based hardware '
    'sensors with a Flutter mobile application. The system enables healthcare workers and users to '
    'track vital signs including heart rate, SpO2, temperature, blood pressure, weight, height, and BMI.'
)

add_heading_with_color(doc, 'Purpose', 2)
doc.add_paragraph(
    'The primary purpose of HealthX is to democratize healthcare monitoring by providing affordable, '
    'accessible, and real-time health data tracking for both individual users and healthcare professionals. '
    'The system bridges the gap between traditional healthcare monitoring and modern IoT technology.'
)

add_heading_with_color(doc, 'Key Features', 2)
features = [
    'Real-time health data monitoring from ESP32 sensors',
    'Dual user system: Health Workers and Regular Users',
    'Patient management and tracking capabilities',
    'Comprehensive health readings storage and history',
    'Offline data persistence with Hive database',
    'Automatic ESP32 connection monitoring',
    'BMI calculation and health metrics analysis',
    'User profile management',
    'Secure authentication system',
    'Weight scale calibration and tare functionality',
    'Blood pressure measurement triggering'
]
for feature in features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_page_break()

# ==================== 2. SYSTEM ARCHITECTURE ====================
add_heading_with_color(doc, '2. System Architecture', 1)

doc.add_paragraph(
    'HealthX follows a three-tier client-server architecture with an additional hardware layer, '
    'ensuring separation of concerns and scalability.'
)

add_heading_with_color(doc, 'Architecture Layers', 2)

arch_table = doc.add_table(rows=5, cols=2)
arch_table.style = 'Light Grid Accent 1'

arch_data = [
    ('Layer', 'Description'),
    ('Presentation Layer', 'Flutter mobile application providing user interface (Dart)'),
    ('Business Logic Layer', 'PHP REST API handling application logic and validation'),
    ('Data Layer', 'MySQL database (server) + Hive (local mobile storage)'),
    ('Hardware Layer', 'ESP32 microcontroller with health monitoring sensors')
]

for i, (layer, desc) in enumerate(arch_data):
    row = arch_table.rows[i]
    row.cells[0].text = layer
    row.cells[1].text = desc

doc.add_paragraph()

add_heading_with_color(doc, 'Communication Flow', 2)
flow_steps = [
    'User authenticates through Flutter app using email and password',
    'App communicates with PHP backend API over HTTP/JSON',
    'Backend validates requests and queries MySQL database',
    'ESP32 device provides sensor data via HTTP endpoints',
    'App stores data locally in Hive for offline access and sync',
    'Real-time connection monitoring ensures hardware availability',
    'Data synchronization occurs when internet connection is available'
]
for i, step in enumerate(flow_steps, 1):
    doc.add_paragraph(f'{i}. {step}', style='List Number')

doc.add_page_break()

# ==================== 3. TECHNOLOGY STACK ====================
add_heading_with_color(doc, '3. Technology Stack', 1)

add_heading_with_color(doc, 'Frontend Technologies', 2)
frontend = [
    'Flutter 3.x - Cross-platform mobile framework',
    'Dart - Programming language for Flutter',
    'Hive 2.x - Local NoSQL database for offline storage',
    'HTTP Package - API communication library',
    'Material Design - UI/UX framework'
]
for tech in frontend:
    doc.add_paragraph(tech, style='List Bullet')

add_heading_with_color(doc, 'Backend Technologies', 2)
backend = [
    'PHP 7.4+ - Server-side scripting language',
    'MySQL 5.7+ - Relational database management system',
    'REST API - Architectural style for web services',
    'JSON - Data interchange format',
    'Apache/Nginx - Web server'
]
for tech in backend:
    doc.add_paragraph(tech, style='List Bullet')

add_heading_with_color(doc, 'Hardware Components', 2)
hardware = [
    'ESP32 - Microcontroller with WiFi capability',
    'MAX30102 - Heart rate and SpO2 sensor',
    'MLX90614 - Non-contact temperature sensor',
    'Pressure sensor - Blood pressure measurement',
    'HX711 + Load cell - Weight measurement',
    'HC-SR04 - Ultrasonic height sensor'
]
for tech in hardware:
    doc.add_paragraph(tech, style='List Bullet')

doc.add_page_break()

# ==================== 4. PROJECT STRUCTURE ====================
add_heading_with_color(doc, '4. Project Structure', 1)

doc.add_paragraph(
    'The Flutter application is organized into a modular structure that promotes maintainability, '
    'scalability, and code reusability.'
)

structure = '''project/
├── assets/
│   ├── images/              # Application images and logo
│   │   └── logo.png
│   └── icon/               # Application icon
│       └── icon.png
│
├── lib/
│   ├── models/             # Data models
│   │   ├── health_reading.dart          # Health reading model
│   │   └── health_reading.g.dart        # Generated Hive adapter
│   │
│   ├── api_config.dart                  # API endpoints configuration
│   ├── api_service.dart                 # API communication service
│   ├── esp32_connection_service.dart    # ESP32 monitoring service
│   │
│   ├── main.dart                        # Application entry point
│   ├── onboarding.dart                  # Onboarding screens
│   ├── login.dart                       # Login screen
│   ├── signup.dart                      # Registration screen
│   │
│   ├── healthworker_dashboard.dart      # Health worker main screen
│   ├── user_dashboard.dart              # User main screen
│   │
│   ├── monitor_screen.dart              # Real-time monitoring
│   ├── parameter_measurement_screen.dart # Individual measurements
│   ├── patient_monitoring_screen.dart   # Patient data monitoring
│   ├── patients_screen.dart             # Patient list management
│   ├── settings_screen.dart             # Health worker settings
│   └── user_settings_screen.dart        # User settings
│
└── pubspec.yaml                         # Dependencies and assets'''

para = doc.add_paragraph(structure)
for run in para.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

doc.add_page_break()

# ==================== 5. CORE COMPONENTS ====================
add_heading_with_color(doc, '5. Core Components', 1)

# 5.1 API Configuration
add_heading_with_color(doc, '5.1 API Configuration (api_config.dart)', 2)
doc.add_paragraph(
    'Centralizes all API endpoint URLs for both the web server and ESP32 device. '
    'This design pattern allows easy configuration changes without modifying multiple files throughout the codebase.'
)

doc.add_paragraph('Configuration Structure:', style='Heading 3')
config_items = [
    'Base URL: http://10.219.51.30/healthx/api (Web server)',
    'ESP32 URL: http://192.168.8.52 (Hardware device)',
    'Authentication endpoints (login, signup)',
    'Patient management endpoints (CRUD operations)',
    'Health readings endpoints (save, retrieve, update, delete)',
    'ESP32 sensor endpoints (individual sensor data retrieval)',
    'Calibration endpoints (weight scale configuration)'
]
for item in config_items:
    doc.add_paragraph(item, style='List Bullet')

# 5.2 API Service
add_heading_with_color(doc, '5.2 API Service (api_service.dart)', 2)
doc.add_paragraph(
    'Provides a centralized interface for all HTTP requests to both the backend API and ESP32 device. '
    'Implements error handling, timeout management, and response parsing for reliable communication.'
)

doc.add_paragraph('Service Methods:', style='Heading 3')

method_table = doc.add_table(rows=1, cols=3)
method_table.style = 'Light Grid Accent 1'

hdr = method_table.rows[0].cells
hdr[0].text = 'Method'
hdr[1].text = 'Type'
hdr[2].text = 'Description'

methods = [
    ('login()', 'Auth', 'User authentication with email/password'),
    ('signup()', 'Auth', 'New user registration'),
    ('getPatients()', 'Patient', 'Retrieve all patient records'),
    ('deletePatient()', 'Patient', 'Remove patient from system'),
    ('getReadings()', 'Reading', 'Fetch health readings by email'),
    ('saveReading()', 'Reading', 'Store new health measurement'),
    ('deleteReading()', 'Reading', 'Remove health record'),
    ('updateReading()', 'Reading', 'Modify existing measurement'),
    ('getEsp32HealthData()', 'ESP32', 'Get all sensor data'),
    ('getWeight()', 'ESP32', 'Individual weight reading'),
    ('getHeight()', 'ESP32', 'Individual height reading'),
    ('getHeartRate()', 'ESP32', 'Heart rate measurement'),
    ('getSpO2()', 'ESP32', 'Blood oxygen level'),
    ('getTemperature()', 'ESP32', 'Body temperature'),
    ('getBloodPressure()', 'ESP32', 'Blood pressure reading'),
    ('startBPMeasurement()', 'ESP32', 'Trigger BP measurement'),
    ('tareWeightScale()', 'ESP32', 'Calibrate weight scale'),
    ('updateCalibrationFactor()', 'ESP32', 'Adjust scale calibration')
]

for method, type_, desc in methods:
    row = method_table.add_row()
    row.cells[0].text = method
    row.cells[1].text = type_
    row.cells[2].text = desc

doc.add_paragraph()

# 5.3 ESP32 Connection Service
add_heading_with_color(doc, '5.3 ESP32 Connection Service (esp32_connection_service.dart)', 2)
doc.add_paragraph(
    'A singleton service that continuously monitors ESP32 device connectivity throughout the application lifecycle. '
    'Runs globally and provides real-time connection status to all screens.'
)

doc.add_paragraph('Key Features:', style='Heading 3')
esp32_features = [
    'Singleton pattern ensures single instance across app',
    'Automatic connection checks every 3 seconds',
    'ValueNotifier for reactive UI updates',
    'Prevents redundant checks with internal flag',
    'Starts automatically on app launch in main.dart',
    'Runs continuously across all screens and navigation',
    'Manual refresh capability for user-triggered checks',
    'Proper resource cleanup on app termination'
]
for feature in esp32_features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_page_break()

# 5.4 Data Models
add_heading_with_color(doc, '5.4 Health Reading Model (models/health_reading.dart)', 2)
doc.add_paragraph(
    'Defines the structure for health reading data with Hive annotations for local storage persistence. '
    'Includes serialization methods for seamless API communication and data synchronization.'
)

doc.add_paragraph('Model Fields:', style='Heading 3')

field_table = doc.add_table(rows=1, cols=3)
field_table.style = 'Light Grid Accent 1'

hdr = field_table.rows[0].cells
hdr[0].text = 'Field'
hdr[1].text = 'Type'
hdr[2].text = 'Description'

fields = [
    ('patientName', 'String', 'Name of the patient'),
    ('userEmail', 'String', 'User email identifier'),
    ('weight', 'double', 'Weight in kilograms'),
    ('height', 'double', 'Height in centimeters'),
    ('bmi', 'double', 'Body Mass Index (calculated)'),
    ('heartRate', 'int', 'Heart rate in beats per minute'),
    ('spo2', 'int', 'Blood oxygen saturation percentage'),
    ('temperature', 'double', 'Body temperature in Celsius'),
    ('systolic', 'int', 'Systolic blood pressure (mmHg)'),
    ('diastolic', 'int', 'Diastolic blood pressure (mmHg)'),
    ('timestamp', 'DateTime', 'Reading date and time'),
    ('synced', 'bool', 'Server synchronization status')
]

for field, type_, desc in fields:
    row = field_table.add_row()
    row.cells[0].text = field
    row.cells[1].text = type_
    row.cells[2].text = desc

doc.add_paragraph()
doc.add_paragraph(
    'The model includes toJson() and fromJson() methods for serialization, '
    'and extends HiveObject for local database integration.'
)

doc.add_page_break()

# ==================== 6. API ENDPOINTS ====================
add_heading_with_color(doc, '6. API Endpoints', 1)

add_heading_with_color(doc, '6.1 Web Server API (PHP Backend)', 2)
doc.add_paragraph('Base URL: http://10.219.51.30/healthx/api')
doc.add_paragraph()

api_table = doc.add_table(rows=1, cols=4)
api_table.style = 'Light Grid Accent 1'

hdr = api_table.rows[0].cells
hdr[0].text = 'Endpoint'
hdr[1].text = 'Method'
hdr[2].text = 'Parameters'
hdr[3].text = 'Description'

endpoints = [
    ('/auth.php', 'POST', 'action, email, password, name', 'Authentication and registration'),
    ('/get_patients.php', 'GET', 'None', 'Retrieve all patients'),
    ('/get_readings.php', 'GET', 'user_email', 'Fetch health readings'),
    ('/save_reading.php', 'POST', 'reading data object', 'Store new reading'),
    ('/delete_patient.php', 'POST', 'patient_id', 'Remove patient'),
    ('/delete_reading.php', 'POST', 'reading_id', 'Delete reading'),
    ('/update_reading.php', 'POST', 'reading_id, data', 'Update reading'),
    ('/update_profile.php', 'POST', 'user data', 'Update user profile')
]

for endpoint, method, params, desc in endpoints:
    row = api_table.add_row()
    row.cells[0].text = endpoint
    row.cells[1].text = method
    row.cells[2].text = params
    row.cells[3].text = desc

doc.add_paragraph()

add_heading_with_color(doc, '6.2 ESP32 Sensor API', 2)
doc.add_paragraph('Base URL: http://192.168.8.52')
doc.add_paragraph()

esp32_table = doc.add_table(rows=1, cols=3)
esp32_table.style = 'Light Grid Accent 1'

hdr = esp32_table.rows[0].cells
hdr[0].text = 'Endpoint'
hdr[1].text = 'Method'
hdr[2].text = 'Returns'

esp32_endpoints = [
    ('/health', 'GET', 'All sensor data (composite)'),
    ('/status', 'GET', 'Device status information'),
    ('/weight', 'GET', 'Current weight reading'),
    ('/height', 'GET', 'Current height reading'),
    ('/heartrate', 'GET', 'Heart rate data'),
    ('/spo2', 'GET', 'Blood oxygen level'),
    ('/temperature', 'GET', 'Body temperature'),
    ('/bloodpressure', 'GET', 'BP systolic/diastolic'),
    ('/tare', 'POST', 'Calibration confirmation'),
    ('/bp/start', 'POST', 'Measurement trigger'),
    ('/config', 'POST', 'Calibration update status')
]

for endpoint, method, returns in esp32_endpoints:
    row = esp32_table.add_row()
    row.cells[0].text = endpoint
    row.cells[1].text = method
    row.cells[2].text = returns

doc.add_page_break()

# ==================== 7. USER INTERFACES ====================
add_heading_with_color(doc, '7. User Interfaces', 1)

add_heading_with_color(doc, '7.1 Authentication Screens', 2)

doc.add_paragraph('Onboarding Screen:', style='Heading 3')
doc.add_paragraph(
    'Three-page carousel introducing users to system features including real-time monitoring, '
    'health worker connectivity, and health management capabilities. Includes skip functionality '
    'and smooth page transitions.'
)

doc.add_paragraph('Login Screen:', style='Heading 3')
login_features = [
    'Email and password validation',
    'Password visibility toggle',
    'User type detection (Health Worker vs Regular User)',
    'Automatic navigation to appropriate dashboard',
    'Error handling with user-friendly messages',
    'Hive local storage integration for session persistence'
]
for feature in login_features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_paragraph('Signup Screen:', style='Heading 3')
signup_features = [
    'Full name, email, and password fields',
    'Password confirmation validation',
    'Real-time input validation',
    'Secure password requirements (min 6 characters)',
    'Automatic return to login after successful registration'
]
for feature in signup_features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_paragraph()

add_heading_with_color(doc, '7.2 Health Worker Dashboard', 2)
doc.add_paragraph(
    'Primary interface for healthcare professionals with three main sections accessible via bottom navigation.'
)

doc.add_paragraph('Monitor Screen:', style='Heading 3')
monitor_features = [
    'Real-time ESP32 connection status indicator',
    'Patient selection dropdown',
    'Live sensor data display for all vital signs',
    'Manual and automatic data refresh',
    'Save measurements to database',
    'Visual feedback for connection status',
    'Weight scale tare functionality'
]
for feature in monitor_features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_paragraph('Patients Screen:', style='Heading 3')
patients_features = [
    'Comprehensive patient list view',
    'Patient data cards with summary information',
    'Delete patient functionality with confirmation',
    'Patient-specific reading history access',
    'Real-time data synchronization',
    'Empty state handling'
]
for feature in patients_features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_paragraph('Settings Screen:', style='Heading 3')
settings_features = [
    'Profile information display',
    'Email and name editing',
    'Password change functionality',
    'Profile picture management',
    'Logout functionality',
    'Data persistence with Hive'
]
for feature in settings_features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_paragraph()

add_heading_with_color(doc, '7.3 User Dashboard', 2)
doc.add_paragraph(
    'Simplified interface for regular users to track their own health data.'
)

user_dashboard_features = [
    'Personal health metrics dashboard',
    'Historical data visualization',
    'Individual parameter measurement screens',
    'Reading history with date filtering',
    'Profile management',
    'Settings and preferences'
]
for feature in user_dashboard_features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_page_break()

# ==================== 8. DATA FLOW ====================
add_heading_with_color(doc, '8. Data Flow', 1)

add_heading_with_color(doc, '8.1 Authentication Flow', 2)
auth_steps = [
    'User enters credentials on Login screen',
    'App sends POST request to /auth.php',
    'Backend validates credentials against MySQL database',
    'On success, user data (id, name, email, user_type) is returned',
    'App saves user data to Hive local storage',
    'User is navigated to appropriate dashboard based on user_type',
    'Session persists until logout'
]
for i, step in enumerate(auth_steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_paragraph()

add_heading_with_color(doc, '8.2 Health Reading Collection Flow', 2)
reading_steps = [
    'Health worker selects patient from dropdown',
    'Worker initiates measurement on Monitor screen',
    'App requests data from ESP32 endpoints',
    'ESP32 sensors collect vital signs data',
    'Data is sent back to app in JSON format',
    'App displays real-time data to health worker',
    'Worker reviews and confirms measurement',
    'App sends reading to PHP backend via /save_reading.php',
    'Backend stores data in MySQL database',
    'Reading is also saved locally in Hive for offline access',
    'Success confirmation displayed to user'
]
for i, step in enumerate(reading_steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_paragraph()

add_heading_with_color(doc, '8.3 ESP32 Connection Monitoring Flow', 2)
monitor_steps = [
    'Esp32ConnectionService starts on app launch',
    'Service checks /health endpoint every 3 seconds',
    'Connection status updates ValueNotifier',
    'All screens listening to ValueNotifier update UI',
    'Visual indicator shows connected/disconnected state',
    'Service runs continuously until app closes'
]
for i, step in enumerate(monitor_steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_page_break()

# ==================== 9. INSTALLATION GUIDE ====================
add_heading_with_color(doc, '9. Installation Guide', 1)

add_heading_with_color(doc, '9.1 Prerequisites', 2)
prerequisites = [
    'Flutter SDK 3.0 or higher',
    'Dart SDK 2.17 or higher',
    'Android Studio / VS Code with Flutter extensions',
    'Android SDK (for Android deployment)',
    'Xcode (for iOS deployment - macOS only)',
    'PHP 7.4+ with MySQL',
    'Web server (Apache/Nginx)',
    'ESP32 development board',
    'Required sensors (MAX30102, MLX90614, HX711, HC-SR04)'
]
for item in prerequisites:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()

add_heading_with_color(doc, '9.2 Flutter App Installation', 2)

installation_code = '''# Clone the repository
git clone <repository-url>
cd healthx

# Install dependencies
flutter pub get

# Generate Hive adapters
flutter packages pub run build_runner build

# Run the app
flutter run'''

para = doc.add_paragraph(installation_code)
for run in para.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

doc.add_paragraph()

add_heading_with_color(doc, '9.3 Backend Setup', 2)

backend_steps = [
    'Install PHP and MySQL on your server',
    'Create a new MySQL database named "healthx"',
    'Import the database schema (create tables for users, patients, readings)',
    'Place PHP API files in web server directory (e.g., /var/www/html/healthx/api)',
    'Configure database connection in PHP files',
    'Update ApiConfig.dart with your server IP address',
    'Test API endpoints using Postman or similar tool'
]
for i, step in enumerate(backend_steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_paragraph()

add_heading_with_color(doc, '9.4 ESP32 Setup', 2)

esp32_steps = [
    'Install Arduino IDE with ESP32 board support',
    'Connect sensors to ESP32 according to pin configuration',
    'Upload firmware code to ESP32',
    'Configure WiFi credentials in ESP32 code',
    'Update esp32Url in ApiConfig.dart with ESP32 IP address',
    'Test sensor readings using serial monitor',
    'Verify HTTP endpoints are accessible'
]
for i, step in enumerate(esp32_steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_page_break()

# ==================== 10. CONFIGURATION ====================
add_heading_with_color(doc, '10. Configuration', 1)

add_heading_with_color(doc, '10.1 API Configuration', 2)
doc.add_paragraph('File: lib/api_config.dart')

config_code = '''// Update these URLs based on your deployment
static const String baseUrl = 'http://YOUR_SERVER_IP/healthx/api';
static const String esp32Url = 'http://YOUR_ESP32_IP';'''

para = doc.add_paragraph(config_code)
for run in para.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

doc.add_paragraph()

add_heading_with_color(doc, '10.2 Color Scheme', 2)
doc.add_paragraph(
    'The application uses a consistent color scheme defined in healthworker_dashboard.dart '
    'and applied throughout the app.'
)

colors = [
    'Primary: #1848A0 (Blue)',
    'Success: #10B981 (Green)',
    'Error: #EF4444 (Red)',
    'Warning: #FDF59E0B (Amber)',
    'Background: #F8FAFC (Light Gray)',
    'Text Dark: #1E293B',
    'Text Light: #64748B'
]
for color in colors:
    doc.add_paragraph(color, style='List Bullet')

doc.add_paragraph()

add_heading_with_color(doc, '10.3 Database Configuration', 2)
doc.add_paragraph(
    'Configure MySQL database connection parameters in your PHP backend files.'
)

db_params = [
    'Database Host: localhost (or your MySQL server address)',
    'Database Name: healthx',
    'Database User: your_mysql_username',
    'Database Password: your_mysql_password',
    'Character Set: utf8mb4'
]
for param in db_params:
    doc.add_paragraph(param, style='List Bullet')

doc.add_page_break()

# ==================== 11. USER GUIDE ====================
add_heading_with_color(doc, '11. User Guide', 1)

add_heading_with_color(doc, '11.1 For Health Workers', 2)

doc.add_paragraph('Getting Started:', style='Heading 3')
hw_start = [
    'Download and install HealthX app',
    'Create account using your professional email',
    'Login with credentials',
    'Navigate to Monitor screen to begin'
]
for i, step in enumerate(hw_start, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_paragraph()
doc.add_paragraph('Taking Measurements:', style='Heading 3')
measurements = [
    'Ensure ESP32 device is powered on (check connection indicator)',
    'Select patient from dropdown menu',
    'Ask patient to position themselves for measurement',
    'Click "Tare Scale" if using weight measurement',
    'Wait for sensor readings to stabilize',
    'Review all vital signs displayed',
    'Click "Save Reading" to store data',
    'Confirm successful save message'
]
for i, step in enumerate(measurements, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_paragraph()
doc.add_paragraph('Managing Patients:', style='Heading 3')
patient_mgmt = [
    'Navigate to Patients screen',
    'View list of all patients with recent readings',
    'Tap patient card to view detailed history',
    'Use delete option to remove patient records (with confirmation)',
    'Data syncs automatically with server'
]
for i, step in enumerate(patient_mgmt, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_paragraph()

add_heading_with_color(doc, '11.2 For Regular Users', 2)

user_guide = [
    'Create account with personal email',
    'Login to access personal dashboard',
    'View your health metrics and history',
    'Take measurements when ESP32 is available',
    'Track trends over time',
    'Update profile information in Settings'
]
for i, step in enumerate(user_guide, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_page_break()

# ==================== 12. TROUBLESHOOTING ====================
add_heading_with_color(doc, '12. Troubleshooting', 1)

add_heading_with_color(doc, '12.1 Common Issues', 2)

issues_table = doc.add_table(rows=1, cols=2)
issues_table.style = 'Light Grid Accent 1'

hdr = issues_table.rows[0].cells
hdr[0].text = 'Problem'
hdr[1].text = 'Solution'

issues = [
    ('ESP32 shows disconnected', 'Check WiFi connection, verify ESP32 IP address in config, restart ESP32 device'),
    ('Login fails', 'Verify internet connection, check email/password, ensure backend server is running'),
    ('Readings not saving', 'Check internet connection, verify backend API is accessible, check server logs'),
    ('Weight readings incorrect', 'Calibrate scale using tare function, check load cell connections'),
    ('App crashes on startup', 'Clear app data, reinstall app, check for Flutter/Dart version compatibility'),
    ('Sensor data not updating', 'Check ESP32 power supply, verify sensor connections, restart device'),
    ('Profile updates not persisting', 'Check Hive database initialization, verify local storage permissions')
]

for problem, solution in issues:
    row = issues_table.add_row()
    row.cells[0].text = problem
    row.cells[1].text = solution

doc.add_paragraph()

add_heading_with_color(doc, '12.2 Debug Mode', 2)
doc.add_paragraph(
    'Enable debug logging by checking console output. The app prints detailed logs for:'
)
debug_items = [
    'API requests and responses',
    'ESP32 connection status changes',
    'Hive database operations',
    'Navigation events',
    'Error messages with stack traces'
]
for item in debug_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ==================== 13. FUTURE ENHANCEMENTS ====================
add_heading_with_color(doc, '13. Future Enhancements', 1)

doc.add_paragraph(
    'Planned features and improvements for future versions of HealthX:'
)

enhancements = [
    'Cloud data backup and synchronization',
    'Multi-language support (internationalization)',
    'Data visualization with charts and graphs',
    'Health alerts and notifications',
    'Integration with wearable devices',
    'Telemedicine consultation features',
    'AI-powered health insights and predictions',
    'Export data to PDF/CSV formats',
    'Family health tracking',
    'Medication reminders',
    'Appointment scheduling with healthcare providers',
    'Integration with hospital management systems',
    'Voice commands for hands-free operation',
    'Biometric authentication (fingerprint/face ID)',
    'Dark mode theme'
]
for enhancement in enhancements:
    doc.add_paragraph(enhancement, style='List Bullet')

doc.add_paragraph()

add_heading_with_color(doc, 'Technical Improvements', 2)
tech_improvements = [
    'Implement GraphQL for more efficient data queries',
    'Add unit and integration tests',
    'Implement CI/CD pipeline',
    'Add code documentation and API docs',
    'Optimize database queries and indexing',
    'Implement caching mechanisms',
    'Add real-time WebSocket communication',
    'Enhance security with OAuth 2.0',
    'Implement data encryption at rest and in transit'
]
for improvement in tech_improvements:
    doc.add_paragraph(improvement, style='List Bullet')

doc.add_page_break()

# ==================== APPENDIX ====================
add_heading_with_color(doc, 'Appendix', 1)

add_heading_with_color(doc, 'A. Color Reference', 2)
doc.add_paragraph('Complete color palette used throughout the application:')

color_ref = '''Primary: #1848A0 (RGB: 24, 72, 160)
Success: #10B981 (RGB: 16, 185, 129)
Error: #EF4444 (RGB: 239, 68, 68)
Warning: #FDF59E0B (RGB: 253, 245, 158, 11)
Background: #F8FAFC (RGB: 248, 250, 252)
White: #FFFFFF (RGB: 255, 255, 255)
Text Dark: #1E293B (RGB: 30, 41, 59)
Text Light: #64748B (RGB: 100, 116, 139)'''

para = doc.add_paragraph(color_ref)
for run in para.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

doc.add_paragraph()

add_heading_with_color(doc, 'B. Dependencies', 2)
doc.add_paragraph('Flutter packages required (pubspec.yaml):')

dependencies = '''dependencies:
  flutter:
    sdk: flutter
  hive: ^2.2.3
  hive_flutter: ^1.1.0
  http: ^0.13.5
  
dev_dependencies:
  flutter_test:
    sdk: flutter
  hive_generator: ^2.0.0
  build_runner: ^2.3.3'''

para = doc.add_paragraph(dependencies)
for run in para.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

doc.add_paragraph()

add_heading_with_color(doc, 'C. Support', 2)
doc.add_paragraph(
    'For technical support, bug reports, or feature requests, please contact the development team '
    'or visit the project repository.'
)

doc.add_paragraph()
doc.add_paragraph('---')
doc.add_paragraph('End of Documentation')

# Save document

output_path = 'C:\\xampp\\htdocs\\healthx\\DESIGN PROJECT PAPER\\HealthX_System_Documentation.docx'
doc.save(output_path)
print(f'Documentation created successfully: {output_path}')